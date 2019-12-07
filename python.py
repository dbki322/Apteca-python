import sys
import datetime
from docx import Document
from ui3 import *
from PyQt5 import QtCore, QtGui, QtWidgets, QtSql
from PyQt5.QtSql import QSqlDatabase, QSqlQuery, QSqlQueryModel
from PyQt5.QtWidgets import (QWidget, QPushButton, QLineEdit,
    QInputDialog, QApplication, QMessageBox)

class MyWin(QtWidgets.QMainWindow):
    def __init__(self, parent=None):
        QtWidgets.QWidget.__init__(self, parent)
        self.ui = Ui_MainWindow()
        self.ui.setupUi(self)
        self.bd()

        # События нажатий на кнопки
        self.ui.pushButton.clicked.connect(self.PB1)
        self.ui.pushButton_2.clicked.connect(self.PB2)
        self.ui.pushButton_3.clicked.connect(self.PB3)
        self.ui.pushButton_4.clicked.connect(self.PB4)    
        self.ui.pushButton_10.clicked.connect(self.GoToMain)
        self.ui.pushButton_11.clicked.connect(self.showDialog)
        self.ui.pushButton_5.clicked.connect(self.Add)
        self.ui.pushButton_6.clicked.connect(self.Delete)
        self.ui.pushButton_7.clicked.connect(self.Edit) 
        self.ui.pushButton_8.clicked.connect(self.DateA)
        self.ui.pushButton_9.clicked.connect(self.DateS)
        self.ui.pushButton_13.clicked.connect(self.expired1)
        self.ui.pushButton_16.clicked.connect(self.expired2)
        self.ui.pushButton_14.clicked.connect(self.Nexpired1)
        self.ui.pushButton_15.clicked.connect(self.Nexpired2)
        self.ui.pushButton_12.clicked.connect(self.ReportS)
        self.ui.pushButton_17.clicked.connect(self.ReportA)

        # Функция вывода данных в таблицы при открытии приложения              
    def bd(self):
        db = QtSql.QSqlDatabase.addDatabase('QSQLITE')
        db.setDatabaseName("db.db")
        db.open()
        qry = QSqlQuery()
        qry = db.exec("SELECT  mId AS 'ID',"
                        "mName AS 'Название',"
                        "mDescription AS 'Описание',"
                        "mPrice AS 'Цена' FROM medications")
        tabmodel = QSqlQueryModel()
        self.ui.tableView_2.setModel(tabmodel)
        tabmodel.setQuery(qry)
        self.ui.tableView_2.setColumnWidth( 2, 310 )
        header = self.ui.tableView_2.horizontalHeader()
        Lheader = self.ui.tableView_2.verticalHeader()
        Lheader.setVisible(False)
        header.setStretchLastSection(True)
        self.ui.tableView_2.show
        db.close()

# Функция добавления в бд
    def Add(self):
        text1, ok1 = QInputDialog.getText(self, 'Добавить',
            'Введите название препарата. ')
        if ok1:
            text2,ok2 = QInputDialog.getText(self, 'Добавить',
            'Введите описание препарата. ')
            if ok2:
                text3, ok3 = QInputDialog.getText(self, 'Добавить',
                'Введите цену препарата. ')
                if ok3:
                    db = QtSql.QSqlDatabase.addDatabase('QSQLITE')
                    db.setDatabaseName("db.db")
                    db.open()
                    queryAdd = QSqlQuery()
                    queryAdd.prepare("INSERT INTO medications (mName, mDescription, mPrice) "
                                    "VALUES (?, ?, ?)")          
                    queryAdd.bindValue(0, text1)
                    queryAdd.bindValue(1, text2)
                    queryAdd.bindValue(2, text3)
                    queryAdd.exec()
                    qry = QSqlQuery()
                    qry = db.exec("SELECT mId AS 'ID',"
                                "mName AS 'Название',"
                                "mDescription AS 'Описание',"
                                "mPrice AS 'Цена' FROM medications")
                    tabmodel = QSqlQueryModel()
                    self.ui.tableView_2.setModel(tabmodel)
                    tabmodel.setQuery(qry)
                    self.ui.tableView_2.setColumnWidth( 2, 310 )
                    header = self.ui.tableView_2.horizontalHeader()
                    Lheader = self.ui.tableView_2.verticalHeader()
                    Lheader.setVisible(False)
                    header.setStretchLastSection(True)
                    self.ui.tableView_2.show
                    db.close()

# Функция Посмотреть кол-во на складе
    def showDialog(self):
        text, ok = QInputDialog.getText(self, 'Количество выбранного препарата на складе',
            'Введите ID препарата, количество которого вы хотите посмотреть:')
        if ok:
            db = QtSql.QSqlDatabase.addDatabase('QSQLITE')
            db.setDatabaseName("db.db")
            db.open()
            query = QSqlQuery()
            query.prepare("SELECT SUM (bCol) FROM batch WHERE b_mId = ( ? ) and bDateS is null ")           
            query.bindValue(0, text)
            query.exec()
            while query.next():
                ff =(query.value(0))
                fff = str(ff)
            if fff =='':
                fff='0'
            query2 = QSqlQuery()
            query2.prepare("SELECT SUM (bCol) FROM batch WHERE b_mId = ( ? ) and bDateS not null")           
            query2.bindValue(0, text)
            query2.exec()
            while query2.next():
                aa =(query2.value(0))
                aaa = str(aa)
            if aaa =='':
                aaa='0'
            x=int(aaa)
            y=int(fff)
            suma= y - x
            if suma<0:
                suma=0
            summa=str(suma)
            SumCol = QMessageBox.information(self, "Количество выбранного препарата на складе","Сейчас на складе следующее количество пачек: " + summa)
            db.close()

# Функция удаления
    def Delete(self):
        dltxt, dlok = QInputDialog.getText(self, 'Удаление',
            'Введите ID препарата, который вы хотите удалить:')
        if dlok:
            db = QtSql.QSqlDatabase.addDatabase('QSQLITE')
            db.setDatabaseName("db.db")
            db.open()
            querydl = QSqlQuery()
            querydl.prepare("DELETE FROM medications WHERE mId= ( ? )")           
            querydl.bindValue(0, dltxt)
            querydl.exec()
            qry = QSqlQuery()
            qry = db.exec("SELECT  mId AS 'ID',"
                        "mName AS 'Название',"
                        "mDescription AS 'Описание',"
                        "mPrice AS 'Цена' FROM medications")
            tabmodel = QSqlQueryModel()
            self.ui.tableView_2.setModel(tabmodel)
            tabmodel.setQuery(qry)
            self.ui.tableView_2.setColumnWidth( 2, 310 )
            header = self.ui.tableView_2.horizontalHeader()
            Lheader = self.ui.tableView_2.verticalHeader()
            Lheader.setVisible(False)
            header.setStretchLastSection(True)
            self.ui.tableView_2.show
            db.close()

# Функция редактирования
    def Edit(self):
        edtxt, edok = QInputDialog.getText(self, 'Редактирование',
            'Введите ID препарата, который вы хотите изменить:')
        if edok:
            text1, ok1 = QInputDialog.getText(self, 'Редактирование',
            'Введите название препарата. ')
            if ok1:
                text2,ok2 = QInputDialog.getText(self, 'Редактирование',
                'Введите описание препарата. ')
                if ok2:
                    text3, ok3 = QInputDialog.getText(self, 'Редактирование',
                    'Введите цену препарата. ')
                    if ok3:
                        db = QtSql.QSqlDatabase.addDatabase('QSQLITE')
                        db.setDatabaseName("db.db")
                        db.open()
                        queryAdd = QSqlQuery()
                        queryAdd.prepare("UPDATE medications SET mName=( ? ), mDescription=( ? ), mPrice=( ? ) WHERE mId=(?) ")  
                        queryAdd.bindValue(0, text1)
                        queryAdd.bindValue(1, text2)
                        queryAdd.bindValue(2, text3)
                        queryAdd.bindValue(3, edtxt) 
                        queryAdd.exec()
                        qry = QSqlQuery()
                        qry = db.exec("SELECT mId AS 'ID',"
                                "mName AS 'Название',"
                                "mDescription AS 'Описание',"
                                "mPrice AS 'Цена' FROM medications")
                        tabmodel = QSqlQueryModel()
                        self.ui.tableView_2.setModel(tabmodel)
                        tabmodel.setQuery(qry)
                        self.ui.tableView_2.setColumnWidth( 2, 310 )
                        header = self.ui.tableView_2.horizontalHeader()
                        Lheader = self.ui.tableView_2.verticalHeader()
                        Lheader.setVisible(False)
                        header.setStretchLastSection(True)
                        self.ui.tableView_2.show
                        db.close()
# Функция ввода данных о приходе
    def DateA(self):
        text1, ok1 = QInputDialog.getText(self, 'Приход препаратов',
            'Введите ID препарата. ')
        if ok1:
            now = datetime.datetime.now()
            text2=now.strftime("%Y-%m-%d")
            text3, ok3 = QInputDialog.getText(self, 'Приход препаратов',
            'Введите количество препаратов в партии. ')
            if ok3:
                text4, ok4 = QInputDialog.getText(self, 'Приход препаратов',
                'Введите срок годности препаратов в партии.(В формате: Год-месяц-день. Пример:2018-12-17) ')
                if ok4:
                        db = QtSql.QSqlDatabase.addDatabase('QSQLITE')
                        db.setDatabaseName("db.db")
                        db.open()
                        queryAdd = QSqlQuery()
                        queryAdd.prepare("INSERT INTO batch (b_mId, bDateA, bCol, bDateEx ) "
                                        "VALUES (?, ?, ?, ?)")         
                        queryAdd.bindValue(0, text1)
                        queryAdd.bindValue(1, text2)
                        queryAdd.bindValue(2, text3)
                        queryAdd.bindValue(3, text4)
                        queryAdd.exec()
                        db.close()
# Функция ввода данных о расходе
    def DateS(self):
        text1, ok1 = QInputDialog.getText(self, 'Расход препаратов',
            'Введите ID препарата. ')
        if ok1:
            now = datetime.datetime.now()
            text2=now.strftime("%Y-%m-%d")
            text3, ok3 = QInputDialog.getText(self, 'Расход препаратов',
            'Введите количество расходованных препаратов в партии. ')
            if ok3:
                db = QtSql.QSqlDatabase.addDatabase('QSQLITE')
                db.setDatabaseName("db.db")
                db.open()
                queryAdd = QSqlQuery()
                queryAdd.prepare("INSERT INTO batch (b_mId, bDateS, bCol) "
                                    "VALUES (?, ?, ?)")         
                queryAdd.bindValue(0, text1)
                queryAdd.bindValue(1, text2)
                queryAdd.bindValue(2, text3)
                queryAdd.exec()
                db.close()
# Функция Списать все просроченнные
    def Woff(self):
        db = QtSql.QSqlDatabase.addDatabase('QSQLITE')
        db.setDatabaseName("db.db")
        db.open()
        queryAdd = QSqlQuery()
        now = datetime.datetime.now()
        NowDate=now.strftime("%Y-%m-%d")
        queryAdd.prepare("S INTO batch (b_mId, bDateS, bCol) "
                            "VALUES (?, ?, ?)")         
        queryAdd.exec()
        db.close()
# Посмотреть где истёк срок годности
    def expired1(self):
        db = QtSql.QSqlDatabase.addDatabase('QSQLITE')
        db.setDatabaseName("db.db")
        db.open()
        qry = QSqlQuery()
        now = datetime.datetime.now()
        Date=now.strftime("%Y-%m-%d")
        DateNow=str(Date)
        qry.prepare("SELECT bDateEx AS 'Срок годности', bId AS 'ID партии', mName AS 'Название' FROM batch, medications WHERE bDateEx < ( ? ) and bDateS is null and b_mId= mId ORDER BY bDateEx ASC")           
        qry.bindValue(0, DateNow)
        qry.exec()
        tabmodel = QSqlQueryModel()
        self.ui.tableView.setModel(tabmodel)
        tabmodel.setQuery(qry)
        #self.ui.tableView.setColumnWidth( 2, 310 )
        header = self.ui.tableView.horizontalHeader()
        Lheader = self.ui.tableView.verticalHeader()
        Lheader.setVisible(False)
        header.setStretchLastSection(True)
        self.ui.tableView.show
        db.close()
    def expired2(self):
        db = QtSql.QSqlDatabase.addDatabase('QSQLITE')
        db.setDatabaseName("db.db")
        db.open()
        qry = QSqlQuery()
        now = datetime.datetime.now()
        Date=now.strftime("%Y-%m-%d")
        DateNow=str(Date)
        qry.prepare("SELECT bDateEx AS 'Срок годности', bId AS 'ID партии', mName AS 'Название' FROM batch, medications WHERE bDateEx < ( ? ) and bDateS is null and b_mId= mId ORDER BY bDateEx DESC")           
        qry.bindValue(0, DateNow)
        qry.exec()
        tabmodel = QSqlQueryModel()
        self.ui.tableView.setModel(tabmodel)
        tabmodel.setQuery(qry)
        #self.ui.tableView.setColumnWidth( 2, 310 )
        header = self.ui.tableView.horizontalHeader()
        Lheader = self.ui.tableView.verticalHeader()
        Lheader.setVisible(False)
        header.setStretchLastSection(True)
        self.ui.tableView.show
        db.close()
#Посмотреть партии лекарст где срок годности не истёк отсортированные по возрастанию и по убыванию
    def Nexpired1(self):
        db = QtSql.QSqlDatabase.addDatabase('QSQLITE')
        db.setDatabaseName("db.db")
        db.open()
        qry = QSqlQuery()
        now = datetime.datetime.now()
        Date=now.strftime("%Y-%m-%d")
        DateNow=str(Date)
        qry.prepare("SELECT bDateEx AS 'Срок годности', bId AS 'ID партии', mName AS 'Название' FROM batch, medications WHERE bDateEx > ( ? ) and bDateS is null and b_mId= mId ORDER BY bDateEx DESC")           
        qry.bindValue(0, DateNow)
        qry.exec()
        tabmodel = QSqlQueryModel()
        self.ui.tableView.setModel(tabmodel)
        tabmodel.setQuery(qry)
        #self.ui.tableView.setColumnWidth( 2, 310 )
        header = self.ui.tableView.horizontalHeader()
        Lheader = self.ui.tableView.verticalHeader()
        Lheader.setVisible(False)
        header.setStretchLastSection(True)
        self.ui.tableView.show
        db.close()
    def Nexpired2(self):
        db = QtSql.QSqlDatabase.addDatabase('QSQLITE')
        db.setDatabaseName("db.db")
        db.open()
        qry = QSqlQuery()
        now = datetime.datetime.now()
        Date=now.strftime("%Y-%m-%d")
        DateNow=str(Date)
        qry.prepare("SELECT bDateEx AS 'Срок годности', bId AS 'ID партии', mName AS 'Название' FROM batch, medications WHERE bDateEx > ( ? ) and bDateS is null and b_mId= mId ORDER BY bDateEx ASC")           
        qry.bindValue(0, DateNow)
        qry.exec()
        tabmodel = QSqlQueryModel()
        self.ui.tableView.setModel(tabmodel)
        tabmodel.setQuery(qry)
        #self.ui.tableView.setColumnWidth( 2, 310 )
        header = self.ui.tableView.horizontalHeader()
        Lheader = self.ui.tableView.verticalHeader()
        Lheader.setVisible(False)
        header.setStretchLastSection(True)
        self.ui.tableView.show
        db.close()

    def ReportA(self):
        Alldoxod = 0
        document = Document()
        document.add_heading('Отчёт по приходу препаратов', 0)
        table = document.add_table(rows=1, cols=6)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Id партии'
        hdr_cells[1].text = 'Количество'
        hdr_cells[2].text = 'Срок годности'
        hdr_cells[3].text = 'Дата прихода'
        hdr_cells[4].text = 'Название'
        hdr_cells[5].text = 'Цена'
        var_d1 = self.ui.dateEdit.date().toPyDate()
        var_d2 = self.ui.dateEdit_2.date().toPyDate()
        dSQLf= var_d1.strftime("%Y-%m-%d")
        dSQLl= var_d2.strftime("%Y-%m-%d")
        db = QtSql.QSqlDatabase.addDatabase('QSQLITE')
        db.setDatabaseName("db.db")
        db.open()
        query = QSqlQuery()
        query.prepare("SELECT bId, bCol, bDateEx, bDateA, mName, mPrice  FROM batch, medications WHERE b_mId = mId AND bDateS IS NULL AND bDateA > ( ? ) AND bDateA < (?) ORDER BY bDateEx ASC")           
        query.bindValue(0, dSQLf)
        query.bindValue(1, dSQLl)
        query.exec()
        while query.next():
            Col1 =str(query.value(0))
            Col2 =(query.value(1))
            Col3 =(query.value(2))
            Col4 =(query.value(3))
            Col5 =(query.value(4))
            Col6 =str(query.value(5))
            doxod = int(Col2)*int(Col6)
            Alldoxod += doxod
            
            row_cells = table.add_row().cells
            row_cells[0].text = Col1
            row_cells[1].text = Col2
            row_cells[2].text = Col3
            row_cells[3].text = Col4
            row_cells[4].text = Col5
            row_cells[5].text = Col6
        Alldoxod2 = str(Alldoxod)
        document.add_paragraph('Было закуплено за промежуток времени: от ' + dSQLf + ' до ' + dSQLl + ' на: ' + Alldoxod2 + ' рублей'  , style='Intense Quote')
        document.save('Отчёт по приходу препаратов.docx')
        db.close()

    def ReportS(self):
        Alldoxod = 0
        document = Document()
        document.add_heading('Отчёт по расходу препаратов', 0)
        table = document.add_table(rows=1, cols=5)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Id партии'
        hdr_cells[1].text = 'Количество'
        hdr_cells[2].text = 'Дата расхода'
        hdr_cells[3].text = 'Название'
        hdr_cells[4].text = 'Цена'
        var_d1 = self.ui.dateEdit.date().toPyDate()
        var_d2 = self.ui.dateEdit_2.date().toPyDate()
        dSQLf= var_d1.strftime("%Y-%m-%d")
        dSQLl= var_d2.strftime("%Y-%m-%d")
        db = QtSql.QSqlDatabase.addDatabase('QSQLITE')
        db.setDatabaseName("db.db")
        db.open()
        query = QSqlQuery()
        query.prepare("SELECT bId, bCol, bDateS, mName, mPrice  FROM batch, medications WHERE b_mId = mId AND bDateA IS NULL AND bDateS > ( ? ) AND bDateS < (?) ORDER BY bDateEx ASC")           
        query.bindValue(0, dSQLf)
        query.bindValue(1, dSQLl)
        query.exec()
        while query.next():
            Col1 =str(query.value(0))
            Col2 =(query.value(1))
            Col3 =(query.value(2))
            Col4 =(query.value(3))
            Col5 =str(query.value(4))
            doxod = int(Col2)*int(Col5)
            Alldoxod += doxod
            row_cells = table.add_row().cells
            row_cells[0].text = Col1
            row_cells[1].text = Col2
            row_cells[2].text = Col3
            row_cells[3].text = Col4
            row_cells[4].text = Col5
        Alldoxod2 = str(Alldoxod)
        document.add_paragraph('Было продано за промежуток времени: от ' + dSQLf + ' до ' + dSQLl + ' на: ' + Alldoxod2 + ' рублей'  , style='Intense Quote')
        document.save('Отчёт по расходу препаратов.docx')
        db.close()
# Переходы между пунктами главного меню
    def PB1(self):
        self.ui.stackedWidget_2.setCurrentIndex(1)
    def PB2(self):
        self.ui.stackedWidget_2.setCurrentIndex(2)
    def PB3(self):
        self.ui.stackedWidget_2.setCurrentIndex(3)
    def PB4(self):
        self.ui.stackedWidget_2.setCurrentIndex(4)
    def GoToMain(self):
        self.ui.stackedWidget_2.setCurrentIndex(0)
       

if __name__=="__main__":
    app = QtWidgets.QApplication(sys.argv)
    myapp = MyWin()
    myapp.show()
    sys.exit(app.exec_())